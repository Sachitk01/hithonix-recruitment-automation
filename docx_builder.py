import io
from docx import Document

def build_riva_report_docx(result):
    """
    Creates a Riva L1 Evaluation Report as DOCX (returns bytes).
    """

    doc = Document()

    doc.add_heading("Riva – L1 Evaluation Report", level=1)
    doc.add_paragraph("Hithonix Solutions Private Limited")

    doc.add_heading("Candidate Summary", level=2)
    doc.add_paragraph(result.match_summary)

    doc.add_heading("Strengths", level=2)
    for s in result.strengths:
        doc.add_paragraph(f"• {s}")

    doc.add_heading("Concerns", level=2)
    for c in result.concerns:
        doc.add_paragraph(f"• {c}")

    doc.add_heading("Red Flags", level=2)
    for r in result.red_flags:
        doc.add_paragraph(f"• {r}")

    doc.add_heading("Communication Signals", level=2)
    doc.add_paragraph(result.communication_signals)

    doc.add_heading("Behavioral Signals", level=2)
    doc.add_paragraph(result.behavioral_signals)

    doc.add_heading("Compensation Alignment", level=2)
    doc.add_paragraph(result.compensation_alignment)

    doc.add_heading("Joining Feasibility", level=2)
    doc.add_paragraph(result.joining_feasibility)

    doc.add_heading("Final Decision", level=2)
    doc.add_paragraph(result.final_decision)

    # Return DOCX bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()


import io
from docx import Document

def build_l2_questionnaire(result):
    """
    Builds a personalized L2 questionnaire using the Riva L1 output.
    """

    doc = Document()

    doc.add_heading("L2 Interview Questionnaire", level=1)
    doc.add_paragraph("Prepared by Riva – AI Recruiter")
    
    doc.add_heading("Core Role Questions", level=2)
    doc.add_paragraph("1. Can you walk us through a recent project where you demonstrated core technical competencies required for this role?")
    doc.add_paragraph("2. Explain a challenging scenario you handled that aligns with responsibilities in this role.")
    doc.add_paragraph("3. Describe your thought process behind key decision-making situations.")

    doc.add_heading("Candidate-Specific Questions", level=2)
    doc.add_paragraph("These questions are tailored based on the L1 evaluation:")
    
    personalized_questions = [
        f"1. You mentioned the following strengths: {', '.join(result.strengths)}. Can you elaborate with real examples?",
        f"2. One concern raised was: {', '.join(result.concerns[:1]) if result.concerns else 'None'}. Please clarify this point.",
        f"3. Can you validate your experience in areas where the transcript indicated possible gaps?",
        f"4. Please explain the alignment between your compensation expectations and the company budget."
    ]
    
    for q in personalized_questions:
        doc.add_paragraph(q)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
